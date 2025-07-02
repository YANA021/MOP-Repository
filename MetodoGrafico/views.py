import json
import numpy as np
import io
from PIL import Image
import base64
from django.http import FileResponse, HttpResponse
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from .tabla_intersecciones import tabla_intersecciones
from .form import ProblemaPLForm
from .models import ProblemaPL
from .solver import resolver_problema_lineal
from .sistema_lineal import pasos_vertices
from django.utils.translation import gettext_lazy as _

def to_native(obj):
   
    if isinstance(obj, np.bool_):
        return bool(obj)
    if isinstance(obj, np.integer):
        return int(obj)
    if isinstance(obj, np.floating):
        return float(obj)
    return obj


def _convert_structure(data):
    
    if isinstance(data, dict):
        return {k: _convert_structure(v) for k, v in data.items()}
    if isinstance(data, list):
        return [_convert_structure(v) for v in data]
    return to_native(data)

def _pasos_objetivo(coef_x1, coef_x2, lista_vertices):
    """Genera pasos de sustitución en la función objetivo."""
    from .utils import format_num
    pasos = []
    for idx, v in enumerate(lista_vertices, start=1):
        x_fmt = format_num(v["x"])
        y_fmt = format_num(v["y"])
        expr = f"{coef_x1}({x_fmt}) + {coef_x2}({y_fmt})"
        pasos.append({
            "punto": f"P{idx}",
            "sustitucion": expr,
            "z": format_num(v["z"]),
            "optimo": v.get("optimo", False),
        })
    return pasos

def home(request):
    """Renderiza la página de inicio."""
    return render(request, "home.html")

def manual(request):
    """Renderiza la página manual."""
    return render(request, "manual.html")

def metodo_grafico(request):
    """Muestra el formulario y procesa el método gráfico."""
    mensaje = ""

    # ------------------------------------------------------------------
    # 1) PETICIÓN POST: procesar formulario
    # ------------------------------------------------------------------
    if request.method == "POST":
        form = ProblemaPLForm(request.POST)
        vertices = []

        if form.is_valid():
            # 1.1  Convertir 'restricciones' (texto JSON) en lista de Python
            restr_raw = form.cleaned_data.get("restricciones", "[]")
            if isinstance(restr_raw, str):
                try:
                    restricciones = json.loads(restr_raw)
                except json.JSONDecodeError:
                    restricciones = []
            else:
                restricciones = restr_raw

            # 1.2  Guardar problema en la BD si el usuario está autenticado
            if request.user.is_authenticated:
                ProblemaPL.objects.create(
                    user=request.user,
                    objetivo=form.cleaned_data["objetivo"],
                    coef_x1=form.cleaned_data["coef_x1"],
                    coef_x2=form.cleaned_data["coef_x2"],
                    restricciones=restricciones,
                )
                mensaje = str(_("Problema guardado correctamente."))
            else:
                mensaje = "Inicia sesión para guardar el problema en tu historial."

            # 1.3  Preparar límites opcionales
            limites = {
                "x1_min": form.cleaned_data.get("x1_min"),
                "x1_max": form.cleaned_data.get("x1_max"),
                "x2_min": form.cleaned_data.get("x2_min"),
                "x2_max": form.cleaned_data.get("x2_max"),
            }

            # 1.4  Resolver con el método gráfico
            resultado = resolver_problema_lineal(
                form.cleaned_data["objetivo"],
                form.cleaned_data["coef_x1"],
                form.cleaned_data["coef_x2"],
                restricciones,
                limites=limites,
            )

            # 1.5  Separar la gráfica del resto del resultado
            grafico = resultado.pop("grafica", "")

            opt_val = resultado.get("z")
            for idx, v in enumerate(resultado.get("vertices", []), start=1):
                vert = {
                    "punto": f"P{idx}",
                    "x1": v["x"],
                    "x2": v["y"],
                    "z": v["z"],
                    "optimo": abs(v["z"] - opt_val) < 1e-6,
                }
                vertices.append({_k: to_native(_v) for _k, _v in vert.items()})

            pasos_objetivo = _pasos_objetivo(
                form.cleaned_data["coef_x1"],
                form.cleaned_data["coef_x2"],
                resultado.get("vertices", [])
            )

            objetivo_text = (
                f"Z = {form.cleaned_data['coef_x1']}x₁ + {form.cleaned_data['coef_x2']}x₂"
            )
            restricciones_fmt = [
                f"{r['coef_x1']}x₁ + {r['coef_x2']}x₂ {r['operador']} {r['valor']}"
             for r in restricciones
            ]

            # cadenas sin subíndices para el cálculo de interceptos
            restr_para_tabla = [
                f"{r['coef_x1']} x1 + {r['coef_x2']} x2 {r['operador']} {r['valor']}"
                for r in restricciones
            ]
            df_tabla, pasos_inter = tabla_intersecciones(restr_para_tabla, incluir_pasos=True)
            tabla = [
                {k: to_native(v) for k, v in fila.items()} for fila in df_tabla.to_dict("records")
            ]
            pasos_sistemas = _convert_structure(pasos_vertices(restr_para_tabla))
            print("PASOS SISTEMA:", pasos_sistemas)

        
            # 1.6  Preparar contexto para la plantilla de resultados
            from .utils import format_num

            for key in ("x", "y", "z"):
                if key in resultado:
                    resultado[key] = format_num(resultado[key])
            context = {
                "form": ProblemaPLForm(),  # formulario limpio para nuevos datos
                "mensaje": mensaje,
                "resultado": resultado,
                "grafico": grafico,
                "objetivo": objetivo_text,
                "restricciones": restricciones_fmt,
                "tabla_inter": tabla,
                "pasos_inter": pasos_inter,
                "pasos_sistemas": pasos_sistemas,
                "vertices": vertices,
                 "pasos_objetivo": pasos_objetivo,
                
            }
            
            return render(request, "resultado.html", context)
    else:
        form = ProblemaPLForm()


    form = ProblemaPLForm()
    return render(request, "nuevo_problema.html", {"form": form, "mensaje": mensaje})


@login_required
def historial_problemas(request):
    """Muestra el historial de problemas del usuario con filtros opcionales."""
    todos = ProblemaPL.objects.filter(user=request.user).order_by("created_at")
    index_map = {p.id: idx + 1 for idx, p in enumerate(todos)}
    qs = todos

    orden = request.GET.get("orden", "old")
    if orden == "new":
        qs = qs.order_by("-created_at")

    obj = request.GET.get("obj")
    if obj in {"max", "min"}:
        qs = qs.filter(objetivo=obj)

    desde = request.GET.get("desde")
    if desde:
        qs = qs.filter(created_at__date__gte=desde)

    hasta = request.GET.get("hasta")
    if hasta:
        qs = qs.filter(created_at__date__lte=hasta)

        qs = list(qs)
    for problema in qs:
        problema.numero = index_map.get(problema.id)

    context = {
        "problemas": qs,
        "orden": orden,
        "obj": obj or "",
        "desde": desde or "",
        "hasta": hasta or "",
    }
    return render(request, "historial.html", context)

@login_required
def ver_problema(request, pk):
    """Muestra el resultado de un problema guardado."""
    problema = get_object_or_404(ProblemaPL, pk=pk, user=request.user)

    resultado = resolver_problema_lineal(
        problema.objetivo,
        problema.coef_x1,
        problema.coef_x2,
        problema.restricciones,
    )

    grafico = resultado.pop("grafica", "")
    vertices = []
    opt_val = resultado.get("z")
    for idx, v in enumerate(resultado.get("vertices", []), start=1):
        vert = {
            "punto": f"P{idx}",
            "x1": v["x"],
            "x2": v["y"],
            "z": v["z"],
            "optimo": abs(v["z"] - opt_val) < 1e-6,
        }
        vertices.append({_k: to_native(_v) for _k, _v in vert.items()})

    pasos_objetivo = _pasos_objetivo(
        problema.coef_x1,
        problema.coef_x2,
        resultado.get("vertices", [])
    )

    objetivo_text = f"Z = {problema.coef_x1}x₁ + {problema.coef_x2}x₂"
    restricciones_fmt = [
        f"{r['coef_x1']}x₁ + {r['coef_x2']}x₂ {r['operador']} {r['valor']}"
        for r in problema.restricciones
    ]

    restr_para_tabla = [
        f"{r['coef_x1']} x1 + {r['coef_x2']} x2 {r['operador']} {r['valor']}"
        for r in problema.restricciones
    ]
    df_tabla, pasos_inter = tabla_intersecciones(restr_para_tabla, incluir_pasos=True)
    tabla = [
        {k: to_native(v) for k, v in fila.items()} for fila in df_tabla.to_dict("records")
    ]
    pasos_sistemas = pasos_vertices(restr_para_tabla)

    context = {
        "form": ProblemaPLForm(),
        "mensaje": "",
        "resultado": resultado,
        "grafico": grafico,
        "objetivo": objetivo_text,
        "restricciones": restricciones_fmt,
        "tabla_inter": tabla,
        "pasos_inter": pasos_inter,
        "pasos_sistemas": pasos_sistemas,
        "vertices": vertices,
        "pasos_objetivo": pasos_objetivo,
        "problema_id": pk,
    }
    return render(request, "resultado.html", context)


def resolver_sistema(request):
    """Resuelve un sistema 2x2 y muestra la gráfica con los pasos."""
    html = ""
    if request.method == "POST":
        from .form import SistemaLinealForm
        from .sistema_lineal import resolver_sistema_pasos

        form = SistemaLinealForm(request.POST)
        if form.is_valid():
            res = resolver_sistema_pasos(
                form.cleaned_data["ecuacion1"],
                form.cleaned_data["ecuacion2"],
                metodo=form.cleaned_data["metodo"],
            )
            html = res["html"]
    else:
        form = SistemaLinealForm()

    context = {
        "form": form,
        "sistema_html": html,
    }
    return render(request, "resultado.html", context)


@login_required
def export_pdf(request, pk):

    # Obtener el problema o devolver 404 si no pertenece al usuario
    problema = get_object_or_404(ProblemaPL, pk=pk, user=request.user)

    try:
        # Recalcular el resultado para obtener la figura actualizada
        resultado = resolver_problema_lineal(
            problema.objetivo,
            problema.coef_x1,
            problema.coef_x2,
            problema.restricciones,
        )
        fig = resultado.get("fig")
        if fig is None:
            raise ValueError("Figura no generada")

       # Exportar la figura a PNG y convertirla a PDF
        png_buffer = io.BytesIO()
        fig.write_image(png_buffer, format="png")
        png_buffer.seek(0)

        img = Image.open(png_buffer)
        buffer = io.BytesIO()
        img.save(buffer, format="PDF")
        buffer.seek(0)
    except Exception:
        # Si ocurre cualquier problema devolver un error interno
        return HttpResponse(
            "Error al generar el PDF", status=500, content_type="text/plain"
        )

    nombre = f"grafico_{pk}.pdf"
    return FileResponse(
        buffer,
        as_attachment=True,
        filename=nombre,
        content_type="application/pdf",
    )

@csrf_exempt
def export_pdf_image(request):
    """Convierte una imagen PNG enviada en el cuerpo en un PDF descargable."""

    if request.method != "POST":
        return HttpResponse(status=405)

    image_data = request.POST.get("image")
    if not image_data:
        try:
            body = json.loads(request.body.decode())
            image_data = body.get("image")
        except Exception:
            image_data = None

    if not image_data:
        return HttpResponse("Sin datos de imagen", status=400, content_type="text/plain")

    if image_data.startswith("data:"):
        image_data = image_data.split(",", 1)[1]

    try:
        img_bytes = base64.b64decode(image_data)
        img = Image.open(io.BytesIO(img_bytes))
        buffer = io.BytesIO()
        img.save(buffer, format="PDF")
        buffer.seek(0)
    except Exception:
        return HttpResponse("Error al generar el PDF", status=500, content_type="text/plain")

    return FileResponse(
        buffer,
        as_attachment=True,
        filename="grafico.pdf",
        content_type="application/pdf",
    )
@login_required
def export_procedure(request, pk, fmt):
    """Exporta todo el procedimiento del método gráfico en PDF, Excel o Word."""

    problema = get_object_or_404(ProblemaPL, pk=pk, user=request.user)

    resultado = resolver_problema_lineal(
        problema.objetivo,
        problema.coef_x1,
        problema.coef_x2,
        problema.restricciones,
    )

    objetivo_text = f"Z = {problema.coef_x1}x₁ + {problema.coef_x2}x₂"
    restr_para_tabla = [
        f"{r['coef_x1']} x1 + {r['coef_x2']} x2 {r['operador']} {r['valor']}"
        for r in problema.restricciones
    ]

    df_tabla, pasos_inter = tabla_intersecciones(restr_para_tabla, incluir_pasos=True)
    pasos_sistemas = pasos_vertices(restr_para_tabla)

    vertices = []
    opt_val = resultado.get("z")
    for idx, v in enumerate(resultado.get("vertices", []), start=1):
        vert = {
            "Punto": f"P{idx}",
            "x1": v["x"],
            "x2": v["y"],
            "z": v["z"],
            "optimo": abs(v["z"] - opt_val) < 1e-6,
        }
        vertices.append(vert)

    pasos_objetivo = _pasos_objetivo(
        problema.coef_x1,
        problema.coef_x2,
        resultado.get("vertices", []),
    )

    if fmt == "excel":
        import pandas as pd
        import io

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            pd.DataFrame(pasos_inter).to_excel(writer, sheet_name="Paso1", index=False)
            df_tabla.to_excel(writer, sheet_name="Paso2", index=False)
            pd.DataFrame(pasos_sistemas).to_excel(writer, sheet_name="Paso3", index=False)
            pd.DataFrame(pasos_objetivo).to_excel(writer, sheet_name="Paso4", index=False)
            pd.DataFrame(vertices).to_excel(writer, sheet_name="Vertices", index=False)
        output.seek(0)
        return FileResponse(
            output,
            as_attachment=True,
            filename=f"metodo_grafico_{pk}.xlsx",
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if fmt == "word":
        import io
        from docx import Document

        doc = Document()
        doc.add_heading("Método Gráfico", level=1)
        doc.add_paragraph(f"Función objetivo: {objetivo_text}")

        doc.add_heading("Paso 1: Intersección con los ejes", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Restricción"
        hdr[1].text = "Sustitución"
        hdr[2].text = "Ecuación"
        hdr[3].text = "Resultado"
        hdr[4].text = "Punto"
        for fila in pasos_inter:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fila["restriccion"])
            row_cells[1].text = str(fila["sustitucion"])
            row_cells[2].text = str(fila["ecuacion"])
            row_cells[3].text = str(fila["resultado"])
            row_cells[4].text = str(fila["punto"])

        doc.add_heading("Paso 2: Tabulación de intersecciones", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Restricción"
        hdr[1].text = "Intercepto x1"
        hdr[2].text = "Intercepto x2"
        hdr[3].text = "Puntos"
        for fila in df_tabla.to_dict("records"):
            row_cells = table.add_row().cells
            row_cells[0].text = str(fila["restriccion"])
            row_cells[1].text = str(fila["intercepto_x1"])
            row_cells[2].text = str(fila["intercepto_x2"])
            row_cells[3].text = str(fila["puntos"])

        if pasos_sistemas:
            doc.add_heading("Paso 3: Intersección de restricciones", level=2)
            for item in pasos_sistemas:
                doc.add_paragraph(f"{item['latex1']} ∩ {item['latex2']}")
                for p in item["pasos"]:
                    doc.add_paragraph(p, style="List Bullet")

        if pasos_objetivo:
            doc.add_heading("Paso 4: Evaluación de la función objetivo", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Punto"
            hdr[1].text = "Sustitución"
            hdr[2].text = "Valor Z"
            for item in pasos_objetivo:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item["punto"])
                row_cells[1].text = f"Z = {item['sustitucion']} = {item['z']}"
                row_cells[2].text = str(item["z"])

        doc.add_heading("Resultado Final", level=1)
        doc.add_paragraph(f"Estado: {resultado['status']}")
        doc.add_paragraph(f"x1: {resultado['x']}")
        doc.add_paragraph(f"x2: {resultado['y']}")
        doc.add_paragraph(f"z: {resultado['z']}")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return FileResponse(
            output,
            as_attachment=True,
            filename=f"metodo_grafico_{pk}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    import io
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet

    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    styles = getSampleStyleSheet()
    elems = []

    elems.append(Paragraph("Método Gráfico", styles["Title"]))
    elems.append(Paragraph(f"Función objetivo: {objetivo_text}", styles["Normal"]))
    elems.append(Spacer(1, 12))

    elems.append(Paragraph("Paso 1: Intersección con los ejes", styles["Heading2"]))
    data = [["Restricción", "Sustitución", "Ecuación", "Resultado", "Punto"]]
    for fila in pasos_inter:
        data.append([
            fila["restriccion"],
            fila["sustitucion"],
            fila["ecuacion"],
            fila["resultado"],
            fila["punto"],
        ])
    tabla = Table(data, repeatRows=1)
    tabla.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.gray),
    ]))
    elems.append(tabla)
    elems.append(Spacer(1, 12))

    elems.append(Paragraph("Paso 2: Tabulación de intersecciones", styles["Heading2"]))
    data = [["Restricción", "Intercepto x1", "Intercepto x2", "Puntos"]]
    for fila in df_tabla.to_dict("records"):
        data.append([
            fila["restriccion"],
            fila["intercepto_x1"],
            fila["intercepto_x2"],
            fila["puntos"],
        ])
    tabla = Table(data, repeatRows=1)
    tabla.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.gray),
    ]))
    elems.append(tabla)
    elems.append(Spacer(1, 12))

    if pasos_sistemas:
        elems.append(Paragraph("Paso 3: Intersección de restricciones", styles["Heading2"]))
        for item in pasos_sistemas:
            elems.append(Paragraph(f"{item['latex1']} ∩ {item['latex2']}", styles["Normal"]))
            for p in item["pasos"]:
                elems.append(Paragraph(p, styles["Bullet"]))
        elems.append(Spacer(1, 12))

    if pasos_objetivo:
        elems.append(Paragraph("Paso 4: Evaluación de la función objetivo", styles["Heading2"]))
        data = [["Punto", "Sustitución", "Valor Z"]]
        for item in pasos_objetivo:
            data.append([
                item["punto"],
                f"Z = {item['sustitucion']} = {item['z']}",
                item["z"],
            ])
        tabla = Table(data, repeatRows=1)
        tabla.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("BACKGROUND", (0, 0), (-1, 0), colors.gray),
        ]))
        elems.append(tabla)
        elems.append(Spacer(1, 12))

    elems.append(Paragraph("Resultado Final", styles["Heading2"]))
    elems.append(Paragraph(f"Estado: {resultado['status']}", styles["Normal"]))
    elems.append(Paragraph(f"x1: {resultado['x']}", styles["Normal"]))
    elems.append(Paragraph(f"x2: {resultado['y']}", styles["Normal"]))
    elems.append(Paragraph(f"z: {resultado['z']}", styles["Normal"]))

    doc.build(elems)
    output.seek(0)
    return FileResponse(
        output,
        as_attachment=True,
        filename=f"metodo_grafico_{pk}.pdf",
        content_type="application/pdf",
    )